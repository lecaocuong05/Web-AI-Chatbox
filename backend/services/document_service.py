from pathlib import Path
from docx import Document
import fitz
import re
import io
import os
import tempfile
import hashlib
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from backend.services.vision_service import describe_image

def split_into_chunks(text, chunk_size=1200, overlap=200):
    """
    Chia chunk nhưng ưu tiên cắt tại:
    1. xuống dòng
    2. dấu chấm
    3. dấu ;
    4. dấu :
    Nếu không có mới cắt cứng.
    """
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    start = 0
    length = len(text)
    while start < length:
        ideal_end = min(start + chunk_size, length)
        if ideal_end >= length:
            piece = text[start:].strip()
            if piece:
                chunks.append(piece)
            break

        split_pos = text.rfind("\n", start, ideal_end)
        if split_pos == -1:
            split_pos = text.rfind(".", start, ideal_end)
        if split_pos == -1:
            split_pos = text.rfind(";", start, ideal_end)
        if split_pos == -1:
            split_pos = text.rfind(":", start, ideal_end)
        if split_pos == -1:
            split_pos = ideal_end
        else:
            split_pos += 1

        chunk = text[start:split_pos].strip()
        if chunk:
            chunks.append(chunk)

        next_start = split_pos - overlap
        if next_start <= start:
            next_start = split_pos
        start = next_start
    return chunks

#===========================================================
TABLE_SEPARATOR_PATTERN = re.compile(r"^\|?[\s\-:|]+\|?$")

def split_markdown_table(text, chunk_size=1200):
    """
    Chia nội dung thành nhiều mảnh, LUÔN LẶP LẠI dòng tiêu đề bảng ở đầu mỗi
    mảnh nếu phát hiện đây là bảng Markdown -> mỗi chunk tự hiểu cột nào là
    cột gì, không phụ thuộc chunk đứng trước.

    Tự động nhận diện: tìm dòng phân cách kiểu "| --- | --- |" bất kỳ đâu
    trong text (không bắt buộc ở đầu, vì VLM đọc ảnh đôi khi thêm vài câu
    dẫn trước bảng). Nếu KHÔNG tìm thấy cấu trúc bảng -> coi là văn bản
    thường, dùng lại split_into_chunks() như bình thường.
    """
    if len(text) <= chunk_size:
        return [text]

    lines = text.split("\n")
    header_idx = None
    for i in range(1, len(lines)):
        if "|" in lines[i - 1] and TABLE_SEPARATOR_PATTERN.match(lines[i].strip()):
            header_idx = i - 1
            break

    if header_idx is None:
        # Không phải bảng Markdown -> chia như văn bản thường
        return split_into_chunks(text, chunk_size)

    preamble = "\n".join(lines[:header_idx]).strip()
    header = lines[header_idx]
    separator = lines[header_idx + 1]
    data_rows = lines[header_idx + 2:]
    header_block = (preamble + "\n" if preamble else "") + header + "\n" + separator

    chunks = []
    current_rows = []
    current_len = len(header_block)

    for row in data_rows:
        row_len = len(row) + 1
        if current_rows and current_len + row_len > chunk_size:
            chunks.append(header_block + "\n" + "\n".join(current_rows))
            current_rows = []
            current_len = len(header_block)
        current_rows.append(row)
        current_len += row_len

    if current_rows:
        chunks.append(header_block + "\n" + "\n".join(current_rows))

    return chunks if chunks else [text]
#===========================================================
def cluster_drawings(rects, distance = 40):
    clusters = []
    for rect in rects:
        merged = False
        for cluster in clusters:
            expand = fitz.Rect(
                cluster.x0 - distance,
                cluster.y0 - distance,
                cluster.x1 + distance,
                cluster.y1 + distance
            )
            if expand.intersects(rect):
                cluster |= rect
                merged = True
                break
        if not merged:
            clusters.append(fitz.Rect(rect))
    return clusters

#===========================================================
def _normalize_table_cell(value):
    """
    Chuẩn hóa text trong cell mà KHÔNG đổi ý nghĩa dữ liệu.

    Đặc biệt sửa trường hợp PDF tách identifier qua nhiều dòng, ví dụ:
    "TINH ID\n_" -> "TINH_ID"
    "TEN CTKM\n_" -> "TEN_CTKM"

    Chỉ áp dụng phép ghép bằng "_" khi cell thực sự có ký tự "_"
    và toàn bộ cell có dạng identifier viết hoa/số/khoảng trắng.
    """
    if value is None:
        return ""

    raw = str(value).strip()

    if not raw:
        return ""

    # Nếu cell có dấu "_" và chỉ gồm ký tự kiểu identifier,
    # coi khoảng trắng/newline là phần bị PDF tách của cùng identifier.
    if (
        "_" in raw
        and re.fullmatch(
            r"[A-ZÀ-Ỹ0-9_\s.\-/]+",
            raw
        )
    ):
        tokens = [
            token
            for token in re.split(
                r"[\s_]+",
                raw
            )
            if token
        ]

        # Chỉ ghép khi có ít nhất 2 token thực.
        if len(tokens) >= 2:
            return "_".join(tokens)

    # Trường hợp thông thường chỉ gộp newline/khoảng trắng.
    return re.sub(
        r"\s+",
        " ",
        raw
    ).strip()


def table_to_markdown(rows):
    if not rows:
        return ""
    clean_rows = []
    for row in rows:
        clean_row = []
        for cell in row:
            if cell is None:
                clean_row.append("")
            else:
                clean_row.append(
                    _normalize_table_cell(cell)
                )
        clean_rows.append(clean_row)

    clean_rows = [
        row for row in clean_rows
        if any(cell != "" for cell in row)
    ]
    if not clean_rows:
        return ""
    
    header = clean_rows[0]
    body = clean_rows[1:]
    col_count = len(header)
    lines = []

    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * col_count) + " |")
    for row in body:
        while len(row) < col_count:
            row.append("")

        row = row[:col_count]
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)

#================== Chuyển bảng thành các câu mô tả ==========
def markdown_table_to_facts(rows):
    """
    Chuyển bảng thành các câu mô tả để embedding tốt hơn.
    Hỗ trợ bảng có ô merge, ô None và ô rỗng.
    """
    if not rows or len(rows) < 2:
        return ""
    # Chuẩn hóa header
    headers = []
    for i, h in enumerate(rows[0]):
        normalized_header = _normalize_table_cell(h)
        if not normalized_header:
            headers.append(f"Cột_{i+1}")
        else:
            headers.append(normalized_header)
    facts = []
    facts.append(
        "Bảng gồm các cột: " + ", ".join(headers)
    )
    for raw_row in rows[1:]:
        # Chuẩn hóa row về list để xử lý an toàn cả list/tuple/None.
        row = list(raw_row or [])

        # Nếu số cột thiếu thì bổ sung; nếu dư thì cắt.
        if len(row) < len(headers):
            row.extend([""] * (len(headers) - len(row)))
        elif len(row) > len(headers):
            row = row[:len(headers)]

        values = []
        for h, v in zip(headers, row):
            if v is None:
                continue
            value = _normalize_table_cell(v)
            if value:
                values.append(f"{h}: {value}")
        if values:
            facts.append("• " + " | ".join(values))
    return "\n".join(facts)
# ============================================================
TABLE_CHUNK_SIZE = 1200

def _clean_table_context(text, max_chars=600):
    """
    Chuẩn hóa ngữ cảnh/caption của bảng.
    Chỉ dùng text thật nằm gần bảng, không tự suy diễn.
    """
    lines = []
    seen = set()

    for raw_line in str(text or "").splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not line:
            continue

        key = line.lower()
        if key in seen:
            continue

        seen.add(key)
        lines.append(line)

    if not lines:
        return ""

    result = ""
    for line in reversed(lines):
        candidate = line if not result else line + "\n" + result
        if len(candidate) > max_chars:
            break
        result = candidate

    return result.strip()

def _combine_table_context(section_title="", local_context=""):
    """
    Ghép tiêu đề section với caption/text ngay trước bảng.
    Không lặp lại cùng một nội dung.
    """
    section_title = _clean_table_context(
        section_title,
        max_chars=250
    )
    local_context = _clean_table_context(
        local_context,
        max_chars=500
    )
    parts = []
    if (
        section_title
        and section_title.lower() != "không có tiêu đề"
    ):
        parts.append(section_title)
    if local_context:
        normalized_local = local_context.lower()
        if (
            not parts
            or parts[-1].lower() not in normalized_local
        ):
            parts.append(local_context)
    return _clean_table_context(
        "\n".join(parts),
        max_chars=600
    )

def _table_title_from_context(
    context,
    fallback="Bảng dữ liệu"
):
    """
    Lấy dòng gần bảng nhất làm title ngắn
    để keyword search/rerank dễ match.
    """
    lines = [
        re.sub(r"\s+", " ", line).strip()
        for line in str(context or "").splitlines()
        if line.strip()
    ]
    if lines:
        title = lines[-1]
        if len(title) > 180:
            title = title[:177].rstrip() + "..."
        return title
    fallback = re.sub(
        r"\s+",
        " ",
        str(fallback or "")
    ).strip()
    return fallback or "Bảng dữ liệu"

def _extract_pdf_table_context(
    page,
    table_bbox,
    boilerplate_lines=None,
    max_blocks=4,
    max_lines=8,
    max_chars=550,
    max_vertical_gap=240
):
    """
    Lấy text nằm NGAY PHÍA TRÊN bảng trên cùng trang
    bằng tọa độ PDF.

    Không hard-code tên bảng hay lĩnh vực.
    """
    boilerplate_lines = boilerplate_lines or set()
    bbox = fitz.Rect(table_bbox)

    candidates = []

    for block in page.get_text("blocks"):
        x0, y0, x1, y1, text, *_ = block
        rect = fitz.Rect(x0, y0, x1, y1)
        if rect.intersects(bbox):
            continue
        if rect.y1 > bbox.y0 + 3:
            continue
        gap = bbox.y0 - rect.y1
        if (
            gap < 0
            or gap > max_vertical_gap
        ):
            continue
        clean_lines = []
        for raw_line in str(text or "").splitlines():
            line = re.sub(
                r"\s+",
                " ",
                raw_line
            ).strip()
            if (
                not line
                or line in boilerplate_lines
            ):
                continue
            clean_lines.append(line)
        if clean_lines:
            candidates.append(
                (
                    gap,
                    rect.y0,
                    rect.x0,
                    "\n".join(clean_lines)
                )
            )
    if not candidates:
        return ""
    nearest = sorted(
        candidates,
        key=lambda x: (
            x[0],
            -x[1]
        )
    )[:max_blocks]

    nearest.sort(
        key=lambda x: (
            x[1],
            x[2]
        )
    )
    lines = []
    for _, _, _, text in nearest:
        lines.extend(
            text.splitlines()
        )
    lines = lines[-max_lines:]
    return _clean_table_context(
        "\n".join(lines),
        max_chars=max_chars
    )

def build_table_content(rows, context=""):
    """
    Tạo nội dung hoàn chỉnh cho một phần của bảng.

    Gồm:
    - context/caption gần bảng;
    - facts để embedding/search tốt;
    - Markdown để Qwen đọc đúng hàng/cột.
    """
    if not rows or len(rows) < 2:
        return ""

    facts = markdown_table_to_facts(rows)
    md_table = table_to_markdown(rows)
    context = _clean_table_context(context)

    if (
        not facts.strip()
        or not md_table.strip()
    ):
        return ""

    context_block = ""

    if context:
        context_block = (
            "Ngữ cảnh bảng:\n"
            + context
            + "\n========================\n"
        )

    return f"""Loại dữ liệu: TABLE
{context_block}{facts}
========================
{md_table}""".strip()


def split_table_rows(
    rows,
    max_chars=TABLE_CHUNK_SIZE,
    context=""
):
    """
    Chia bảng lớn theo HÀNG.
    Luôn giữ header + context ở đầu mỗi chunk.
    Không cắt một hàng ra làm đôi.
    """
    if not rows or len(rows) < 2:
        return []

    header = list(rows[0] or [])
    body = rows[1:]

    result = []
    current_rows = [header]

    for raw_row in body:
        row = list(raw_row or [])

        candidate = current_rows + [row]
        candidate_text = build_table_content(
            candidate,
            context=context
        )

        if (
            len(candidate_text) > max_chars
            and len(current_rows) > 1
        ):
            result.append(current_rows)
            current_rows = [header, row]
        else:
            current_rows.append(row)

    if len(current_rows) > 1:
        result.append(current_rows)

    return result

# ================= Hàm vision chung =========================
def vision_to_chunk(image_bytes, title):
    """
    image_bytes: bytes ảnh
    title      : tiêu đề chunk
    """
    chunks = []
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(
            suffix=".png",
            delete=False
        ) as temp:
            temp.write(image_bytes)
            temp_path = temp.name
        description = describe_image(temp_path)

        if description.strip() and "không chứa thông tin có giá trị tra cứu" not in description.lower():
            pieces = split_markdown_table(description)
            for piece in pieces:
                chunks.append({
                    "title": title,
                    "content": piece
                })
    except Exception as e:
        print("Vision:", e)
    finally:
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)
    return chunks 

#===========================================================
def iter_docx_blocks(document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, document)

#=============================================================
TABLE_PATTERN = re.compile(r'^\|.*\|$')
def is_markdown_table(lines, index):
    """Kiểm tra xem từ dòng index có bắt đầu một bảng Markdown hay không."""
    if index + 1 >= len(lines):
        return False

    line1 = lines[index].strip()
    line2 = lines[index + 1].strip()

    return bool(
        TABLE_PATTERN.match(line1)
        and re.match(r'^\|?[\-\:\s|]+\|?$', line2)
    )

def collection_markdown_table(lines, start):
    """
    Thu toàn bộ bảng Markdown thành 1 block"""
    table = []
    i = start 
    while i < len(lines):
        line = lines[i].rstrip()
        if TABLE_PATTERN.match(line):
            table.append(line)
            i += 1
        else:
            break
    return "\n".join(table), i
#===================== TXT ===================================
def read_txt(file_path):
    with open(file_path, "r",encoding="utf-8") as f:
        lines = f.readlines()

    chunks = []
    current_title = "Không có tiêu đề"
    current_content = []
    global_chunk_index = 1

    def save_chunk():
        nonlocal current_content, global_chunk_index
        if current_content:
            text = "\n".join(current_content).strip()
            pieces = split_into_chunks(text)
            for piece in pieces:
                chunks.append({
                    "title": current_title,
                    "content": piece,
                    "type": "text",
                    "page": None,
                    "headers": [],
                    "chunk_index": global_chunk_index,
                    "total_chunks": 0
                })
                global_chunk_index += 1
            current_content = []

    for line in lines:
        text = line.strip()
        if not text:
            continue
        if(
            text.startswith("#")
            or re.match(r"^\d+\.", text)
            or text.upper().startswith("CHƯƠNG")
            or text.upper().startswith("PHẦN")
            or text.upper().startswith("MỤC")
            or re.match(r"^=+$", text)
            or re.match(r"^-+$", text)
        ):
            save_chunk()
            current_title = text.lstrip("#").strip()
        else:
            current_content.append(text)
    save_chunk()
    total = len(chunks)
    for c in chunks:
        c["total_chunks"] = total
    return chunks

#===================== DOCX ===================================
def read_docx(file_path):
    doc = Document(file_path)
    relationship_images = {}
    total_images = len([
        rel for rel in doc.part.rels.values()
        if "image" in rel.target_ref
    ])
    current_image = 0
    print(f"Tổng số ảnh cần xử lý: {total_images}")
    vision_cache = {}
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            relationship_images[rel.rId] = rel.target_part.blob
    chunks = []
    h1 = ""
    h2 = ""
    h3 = ""
    current_content = []
    recent_context_lines = []
    global_chunk_index = 1
    vision_chunk_index = VISION_CHUNK_INDEX_START
    MIN_IMAGE_SIZE_BYTES = 8_000  
    def current_title():
        title = " > ".join(
            x for x in [h1, h2, h3] if x
        )
        return title if title else "Không có tiêu đề"
    def save_chunk():
        nonlocal current_content
        nonlocal global_chunk_index
        if current_content:
            text = "\n".join(current_content).strip()
            pieces = split_into_chunks(text)
            for piece in pieces:
                chunks.append({
                    "title": current_title(),
                    "content": piece,
                    "type": "text",
                    "page": None,
                    "headers": [],
                    "chunk_index": global_chunk_index,
                    "total_chunks": 0
                })
                global_chunk_index += 1
            current_content = []

    for block in iter_docx_blocks(doc):
        if isinstance(block, DocxParagraph):
            text = block.text.strip()
            if text:
                style = block.style.name
                if style.startswith("Heading 1"):
                    save_chunk()
                    h1 = text
                    h2 = ""
                    h3 = ""
                    recent_context_lines = []
                elif style.startswith("Heading 2"):
                    save_chunk()
                    h2 = text
                    h3 = ""
                    recent_context_lines = []
                elif style.startswith("Heading 3"):
                    save_chunk()
                    h3 = text
                    recent_context_lines = []
                else:
                    current_content.append(text)
                    recent_context_lines.append(text)
                    if len(recent_context_lines) > 8:
                        recent_context_lines = recent_context_lines[-8:]
            try:
                for run in block.runs:
                    drawing = run._element.xpath(
                        ".//*[local-name()='blip']"
                    )
                    for item in drawing:
                        r_embed = item.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                        )
                        if (
                            r_embed
                            and
                            r_embed in relationship_images
                        ):
                            image_bytes = relationship_images[r_embed]
                            image_hash = hashlib.md5(image_bytes).hexdigest()
                            if len(image_bytes) < MIN_IMAGE_SIZE_BYTES:
                                continue  
                            if image_hash in vision_cache:
                                current_image += 1
                                print(
                                    f"Ảnh {current_image}/{total_images}"
                                    "(đã cache)"
                                )
                                vision_chunks = vision_cache[image_hash]
                            else:
                                current_image += 1
                                print(
                                    f"Đang xử lý ảnh"
                                    f"{current_image}/{total_images}"
                                )
                                vision_chunks = vision_to_chunk(
                                    image_bytes,
                                    current_title() + " - Vision"
                                )
                                print(
                                    f"Xong ảnh"
                                    f"{current_image}/{total_images}"
                                )
                                vision_cache[image_hash] = vision_chunks
                            for vc in vision_chunks:
                                chunks.append({
                                    "title": vc["title"],
                                    "content": vc["content"],
                                    "type": "vision",
                                    "headers": [],
                                    "page": None,
                                    "chunk_index": vision_chunk_index,
                                    "total_chunks": 0
                                })
                                vision_chunk_index += 1
            except Exception as e:
                print("Vision DOCX:", e)
        elif isinstance(block, DocxTable):
            local_context = "\n".join(
                recent_context_lines[-6:]
            )
            table_context = _combine_table_context(
                current_title(),
                local_context
            )
            table_context_title = _table_title_from_context(
                table_context,
                fallback=current_title()
            )
            save_chunk()

            rows = []
            for row in block.rows:
                rows.append(
                    [
                        cell.text.strip()
                        for cell in row.cells
                    ]
                )
            if not rows or len(rows) < 2:
                recent_context_lines = []
                continue
            headers = [
                str(h).strip()
                if h is not None
                else ""
                for h in rows[0]
            ]
            clean_headers = [
                h
                for h in headers
                if h
            ]

            table_parts = split_table_rows(
                rows,
                context=table_context
            )
            for part_index, part_rows in enumerate(
                table_parts,
                start=1
            ):
                table_content = build_table_content(
                    part_rows,
                    context=table_context
                )
                if not table_content:
                    continue
                if len(table_parts) == 1:
                    table_title = (
                        table_context_title
                        + " (Bảng)"
                    )
                else:
                    table_title = (
                        table_context_title
                        + f" (Bảng - Phần "
                        f"{part_index}/"
                        f"{len(table_parts)})"
                    )

                chunks.append({
                    "title": table_title,
                    "content": table_content,
                    "type": "table",
                    "page": None,
                    "headers": clean_headers,
                    "chunk_index": global_chunk_index,
                    "total_chunks": 0
                })
                global_chunk_index += 1

            recent_context_lines = []

    save_chunk()
    total = len(chunks)
    for c in chunks:
        c["total_chunks"] = total
    return chunks
#=============================================================
def _detect_boilerplate_lines(pdf, threshold_ratio=0.4):
    """
    Quét toàn bộ file để tìm các dòng lặp lại ở phần lớn số trang
    (thường là header/footer như "Phiên bản x.xx Tài liệu ..." hoặc số trang).
    Các dòng này sẽ bị loại khỏi nội dung để không làm nhiễu chunk và retrieval.
    """
    num_pages = len(pdf)
    line_page_count = {}
 
    for page in pdf:
        text = page.get_text()
        seen_this_page = set()
        for line in text.split("\n"):
            line = line.strip()
            if not line or line in seen_this_page:
                continue
            seen_this_page.add(line)
            line_page_count[line] = line_page_count.get(line, 0) + 1
 
    threshold = max(2, int(num_pages * threshold_ratio))
    return {line for line, count in line_page_count.items() if count >= threshold}

#===================== PDF ===================================
FIGURE_CAPTION_PATTERN = re.compile(r"(Hình|Sơ đồ|Biểu đồ)\s+\d+\s*[:.]", re.IGNORECASE)
VISION_CHUNK_INDEX_START = 100_000
DRAWING_COUNT_THRESHOLD = 5 
MIN_IMAGE_SIZE_BYTES_PDF = 8_000

def _non_empty_table_cells(row):
    """
    Lấy các cell có dữ liệu theo đúng thứ tự.
    Dùng để so sánh cấu trúc bảng giữa các trang.
    """
    values = []
    for cell in list(row or []):
        value = _normalize_table_cell(cell)
        if value:
            values.append(value)
    return values

def _header_value_key(value):
    """
    Chuẩn hóa nhẹ header để so sánh.
    Không phụ thuộc hoa/thường hoặc nhiều khoảng trắng.
    """
    return re.sub(
        r"\s+",
        " ",
        str(value or "").strip().lower()
    )

def _is_pdf_table_continuation(
        rows,
        table_bbox,
        page_height,
        page_num,
        previous_tail_table,
        local_context = ""
):
    """
    Xác định bảng ở đầu trang hiện tại có phải phần tiếp nối của bảng cuối trang trước hay không.
    không hard-code tên bảng
    """
    if not previous_tail_table:
        return False
    # Bảng phải đúng trang kế tiếp.
    if previous_tail_table.get("page") != page_num - 1:
        return False
    # Bảng trước phải kết thúc gần cuối trang.
    if previous_tail_table.get("bottom_ratio", 0) < 0.70:
        return False
    
    bbox = fitz.Rect(table_bbox)
    # bảng hiện tại phải bắt đầu gần đầu trang
    if bbox.y0 > page_height * 0.35:
        return False
    previous_headers = previous_tail_table.get(
        "headers",
        []
    )
    if not previous_headers or not rows:
        return False
    first_values = _non_empty_table_cells(
        rows[0]
    )

    #số cột logic phải giống nhau
    if len(first_values) != len(previous_headers):
        return False
    current_keys = [
        _header_value_key(x)
        for x in first_values
    ]
    previous_keys = [
        _header_value_key(x)
        for x in previous_headers
    ]
    if current_keys == previous_keys:
        return False
    return True


def read_pdf(file_path):
    pdf = fitz.open(file_path)
    boilerplate_lines = _detect_boilerplate_lines(pdf)
    chunks = []
    global_chunk_index = 1
    vision_chunk_index = VISION_CHUNK_INDEX_START
    title = "Không có tiêu đề"
    content = []
    total_pages = len(pdf)
    processed_image_hashes = set()
    previous_tail_table = None

    def save_chunk(current_page):
        nonlocal content, global_chunk_index
        if content:
            text = "\n".join(content).strip()
            pieces = split_into_chunks(text)
            for piece in pieces:
                chunks.append({
                    "title": title,
                    "content": piece,
                    "type": "text",
                    "page": current_page,
                    "headers": [],
                    "chunk_index": global_chunk_index,
                    "total_chunks": 0
                })
                global_chunk_index += 1
            content = []
    for page_num, page in enumerate(pdf, start=1):
        print(f"Đang xử lý trang {page_num}/{total_pages}...")
        table_bboxes = []
        current_tail_table = None
        page_text = page.get_text()
        has_real_text = len(page_text.strip()) >= 30  # ngưỡng nhỏ: gần như rỗng -> nghi là trang scan/ảnh

        try:
            found_tables = page.find_tables()
        except Exception:
            found_tables = None

        has_table = bool(found_tables and found_tables.tables)
        has_embedded_image = len(page.get_images(full=True)) > 0
        need_vision = False
        new_embedded_image_bytes = []  # giữ lại ảnh nhúng MỚI để gửi thẳng cho VLM (không cần render cả trang)
        if has_embedded_image:
            for img in page.get_images(full=True):
                xref = img[0]
                try:
                    image_dict = pdf.extract_image(xref)
                    image_bytes = image_dict["image"]
                    if len(image_bytes) < MIN_IMAGE_SIZE_BYTES_PDF:
                        continue
                    image_hash = hashlib.md5(image_bytes).hexdigest()
                    if image_hash not in processed_image_hashes:
                        processed_image_hashes.add(image_hash)
                        need_vision = True
                        new_embedded_image_bytes.append(image_bytes)
                except Exception as e:
                    print(f"Lỗi extract image: {e}")
                    continue
        has_figure_caption = bool(FIGURE_CAPTION_PATTERN.search(page_text))
        drawings = page.get_drawings()
        drawing_count = len(drawings)
        has_vector_diagram = (drawing_count >= DRAWING_COUNT_THRESHOLD) and not has_table
        print(f"   [debug] Trang {page_num}: drawing_count={drawing_count}, "
              f"has_table={has_table}, has_embedded_image={has_embedded_image}, "
              f"has_figure_caption={has_figure_caption}, has_real_text={has_real_text}")
        vision_chunks = []
        if(need_vision or has_figure_caption or has_vector_diagram):
            reason = []
            if need_vision:
                reason.append("có ảnh nhúng mới")
            if has_figure_caption:
                reason.append("có caption Hình/Sơ đồ")
            if has_vector_diagram:
                reason.append(f"nghi có sơ đồ vector ({drawing_count} đối tượng)")
            if not has_real_text:
                print(f"   -> Trang {page_num} {', '.join(reason)}, KHÔNG có text thật "
                      f"(nghi scan) -> gửi nguyên trang cho VLM...")
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                vision_chunks = vision_to_chunk(img_bytes, f"Trang {page_num} - Vision")
                del pix
                del img_bytes
            else:
                print(f"   -> Trang {page_num} {', '.join(reason)}, có text thật "
                      f"-> chỉ gửi vùng ảnh/sơ đồ cho VLM (không gửi cả trang)...")
                # Ưu tiên 1: có ảnh nhúng thật MỚI -> gửi thẳng ảnh đó
                for image_bytes in new_embedded_image_bytes:
                    vc = vision_to_chunk(image_bytes, f"Trang {page_num} - Vision")
                    vision_chunks.extend(vc)
                # Ưu tiên 2: nghi có sơ đồ vector -> cắt đúng vùng chứa các
                # đối tượng vector đó rồi mới render (không render cả trang)
                if has_vector_diagram and drawings:
                    rects = [
                        fitz.Rect(d["rect"])
                        for d in drawings
                        if d.get("rect")
                    ]
                    if rects:
                        x0 = min(r.x0 for r in rects) - 10
                        y0 = min(r.y0 for r in rects) - 10
                        x1 = max(r.x1 for r in rects) + 10
                        y1 = max(r.y1 for r in rects) + 10
                        page_rect = page.rect
                        clip = fitz.Rect(
                            max(x0, 0),
                            max(y0, 0),
                            min(x1, page_rect.width),
                            min(y1, page_rect.height)
                        )
                        if clip.width >= 120 and clip.height >= 120:
                            ratio = clip.width / clip.height
                            if ratio <= 8:
                                if ratio >= 0.12:
                                    page_area = page.rect.get_area()
                                    if clip.get_area() <= page_area * 0.65:
                                        pix = page.get_pixmap(
                                            dpi=300,
                                            clip=clip
                                        )
                                        img_bytes = pix.tobytes("png")
                                        vc = vision_to_chunk(
                                            img_bytes,
                                            f"Trang {page_num} - Vision (sơ đồ)"
                                        )
                                        vision_chunks.extend(vc)
                                        del pix
                                        del img_bytes
        if has_table:
            for index, table in enumerate(
                found_tables.tables,
                start=1
            ):
                rows = table.extract()
                table_bboxes.append(table.bbox)
                if not rows or len(rows) < 2:
                    continue
                local_table_context = _extract_pdf_table_context(
                    page,
                    table.bbox,
                    boilerplate_lines
                )
                is_continuation = _is_pdf_table_continuation(
                    rows=rows,
                    table_bbox=table.bbox,
                    page_height=page.rect.height,
                    page_num=page_num,
                    previous_tail_table=previous_tail_table,
                    local_context=local_table_context
                )
                if is_continuation:
                    inherited_headers = list(
                        previous_tail_table["headers"]
                    )
                    print(
                        f"   [table] Trang {page_num}: "
                        f"phát hiện bảng nối tiếp trang "
                        f"{page_num - 1}"
                    )
                    print(
                        "   [table] Kế thừa headers:",
                        inherited_headers
                    )
                    rows = [inherited_headers] + rows

                if is_continuation:
                    table_context = previous_tail_table.get(
                        "context",
                        ""
                    )
                    if not table_context:
                        table_context = _combine_table_context(
                            title,
                            local_table_context
                        )
                else:
                    table_context = _combine_table_context(
                        title,
                        local_table_context
                    )
                table_context_title = _table_title_from_context(
                    table_context,
                    fallback=title
                )
                headers = [
                    _normalize_table_cell(h)
                    for h in rows[0]
                ]
                clean_headers = [
                    h
                    for h in headers
                    if h
                ]
                # Không có header hợp lệ thì bỏ
                if not clean_headers:
                    continue

                # CHIA BẢNG THÀNH CÁC CHUNK
                table_parts = split_table_rows(
                    rows,
                    context=table_context
                )
                if not table_parts:
                    continue

                # TẠO TABLE CHUNK
                for part_index, part_rows in enumerate(
                    table_parts,
                    start=1
                ):
                    table_content = build_table_content(
                        part_rows,
                        context=table_context
                    )
                    if not table_content:
                        continue
                    if len(table_parts) == 1:
                        table_title = (
                            f"{table_context_title} "
                            f"(Bảng {index} "
                            f"- Trang {page_num})"
                        )
                    else:
                        table_title = (
                            f"{table_context_title} "
                            f"(Bảng {index} - Phần "
                            f"{part_index}/"
                            f"{len(table_parts)} "
                            f"- Trang {page_num})"
                        )
                    chunks.append({
                        "title": table_title,
                        "content": table_content,
                        "type": "table",
                        "page": page_num,
                        "headers": clean_headers,
                        "chunk_index": global_chunk_index,
                        "total_chunks": 0
                    })
                    global_chunk_index += 1
                bbox = fitz.Rect(
                    table.bbox
                )
                bottom_ratio = (
                    bbox.y1
                    / page.rect.height
                )
                if (
                    len(clean_headers) >= 2
                    and bottom_ratio >= 0.70
                ):
                    current_tail_table = {
                        "page": page_num,
                        "headers": list(
                            clean_headers
                        ),
                        "context": table_context,
                        "title": table_context_title,
                        "bottom_ratio": bottom_ratio
                    }
                    print(
                        f"   [table] Trang {page_num}: "
                        f"ghi nhớ bảng cuối trang"
                    )
                    print(
                        "   [table] Headers:",
                        clean_headers
                    )
        blocks = page.get_text("blocks")
        for block in blocks:
            x0, y0, x1, y1, text, *_=block
            block_rect = fitz.Rect(x0, y0, x1, y1)
            inside_table = any(
                block_rect.intersects(fitz.Rect(bbox)) for bbox in table_bboxes
            )
            if inside_table:
                continue

            for line in text.split("\n"):
                line = line.strip()
                if not line or line in boilerplate_lines:
                    continue
                words = line.split()
                if (
                    len(words) <= 10
                    and "." not in line
                    and (
                        line.isupper()
                        or line.endswith(":")
                    )
                ):
                    save_chunk(page_num)
                    title = line
                else:
                    content.append(line)
        for vc in vision_chunks:
            chunks.append({
                "title": vc["title"],
                "content": vc["content"],
                "type": "vision",
                "page": page_num,
                "headers": [],
                "chunk_index": vision_chunk_index,
                "total_chunks": 0
            })
            vision_chunk_index += 1
        previous_tail_table = current_tail_table

    save_chunk(total_pages)
    total = len(chunks)
    for c in chunks:
        c["total_chunks"] = total
    pdf.close()
    return chunks